"""
Resume tailoring via Claude API.
Extracts text from PDF, sends to Claude with job description, returns tailored content.
"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import Optional

import anthropic
from pdfminer.high_level import extract_text
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console
from rich.panel import Panel

console = Console()


def extract_resume_text(pdf_path: str) -> str:
    """Extract plain text from a PDF resume."""
    path = Path(pdf_path).expanduser()
    if not path.exists():
        raise FileNotFoundError(f"Resume not found: {path}")

    console.print(f"[dim]Extracting text from {path.name}...[/dim]")
    text = extract_text(str(path))

    # Clean up excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


def tailor_resume(
    resume_text: str,
    job_description: str,
    company: str,
    role: str,
    config: dict,
) -> dict:
    """
    Send resume + job description to Claude, get tailored resume back.
    Returns dict with tailored_text, score, keywords, summary.
    """
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    prompt = f"""You are an expert resume writer and ATS optimization specialist.

I need you to tailor this resume for a specific job application.

## TARGET JOB
Company: {company}
Role: {role}

## JOB DESCRIPTION
{job_description}

## CURRENT RESUME
{resume_text}

## YOUR TASKS

1. **Tailor the resume** — rewrite it to maximize fit for this specific role:
   - Mirror the exact language and keywords from the job description (ATS optimization)
   - Reorder bullet points to lead with most relevant accomplishments
   - Adjust the professional summary to target this role specifically
   - Quantify any unquantified bullets where reasonable estimates are possible
   - De-emphasize skills/experience irrelevant to this role
   - Keep all facts true — do not invent experience

2. **Score the match** (honest assessment):
   - skills_match: 1-10
   - experience_match: 1-10
   - industry_match: 1-10
   - overall: 1-10
   - gaps: list any significant missing requirements

3. **Extract ATS keywords** found in job description that should be in resume

Return your response as a JSON object with this exact structure:
{{
  "tailored_resume": "FULL RESUME TEXT IN MARKDOWN FORMAT",
  "professional_summary": "2-3 sentence targeted summary for this role",
  "scores": {{
    "skills_match": 8,
    "experience_match": 7,
    "industry_match": 9,
    "overall": 8
  }},
  "gaps": ["Missing: specific certification", "Limited: cloud infrastructure experience"],
  "ats_keywords_added": ["keyword1", "keyword2"],
  "cover_letter_hook": "One compelling sentence to open a cover letter for this role"
}}

Return ONLY the JSON object, no other text."""

    console.print("[dim]Sending to Claude for tailoring...[/dim]")

    model = config.get("claude", {}).get("model", "claude-sonnet-4-6")
    max_tokens = config.get("claude", {}).get("max_tokens", 4096)

    message = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}]
    )

    import json
    response_text = message.content[0].text.strip()

    # Strip markdown code fences if present
    if response_text.startswith("```"):
        response_text = re.sub(r'^```(?:json)?\n', '', response_text)
        response_text = re.sub(r'\n```$', '', response_text)

    result = json.loads(response_text)
    return result


def save_tailored_resume(
    tailored_text: str,
    company: str,
    role: str,
    output_dir: str,
) -> str:
    """Save tailored resume as a .docx file. Returns the file path."""
    output_path = Path(output_dir).expanduser()
    output_path.mkdir(parents=True, exist_ok=True)

    # Sanitize filename
    safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
    safe_role = re.sub(r'[^\w\s-]', '', role).strip().replace(' ', '_')
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"{safe_company}_{safe_role}_{date_str}.docx"
    filepath = output_path / filename

    doc = Document()

    # Style the document
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Parse and write markdown-ish content
    lines = tailored_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
        else:
            doc.add_paragraph(line)

    doc.save(str(filepath))
    console.print(f"[green]Saved tailored resume:[/green] {filepath}")
    return str(filepath)


def generate_cover_letter(
    hook: str,
    company: str,
    role: str,
    config: dict,
) -> str:
    """Generate a cover letter from config template + Claude-generated hook."""
    template = config.get("form_answers", {}).get("cover_letter_template", "")
    personal = config.get("personal", {})

    cover = template.format(
        role=role,
        company=company,
        years_experience=personal.get("years_experience", "several"),
        relevant_skills="software development and engineering",
        tailored_paragraph=hook,
        full_name=personal.get("full_name", ""),
    )
    return cover
